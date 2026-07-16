from pathlib import Path
from loguru import logger
from scraper.detail_scraper import VBDocument

def convert_to_docx(vb: VBDocument, out_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("Missing dependencies. Run: pip install python-docx pypandoc htmldocx beautifulsoup4") from exc

    file_path = out_dir / "noi_dung.docx"
    html_str = getattr(vb, "content_html", "")
    
    # Check document size to prevent freezing on massive tables (e.g. 12MB tariffs)
    is_massive_doc = len(html_str) > 500000

    # 0. FALLBACK DOWNLOADS FOR DOCUMENTS WITHOUT HTML
    is_empty_html = False
    if not html_str or not html_str.strip():
        is_empty_html = True
    else:
        # Some API endpoints return empty HTML boilerplate
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_str, 'html.parser')
        body = soup.find('body')
        if not body or not body.text.strip():
            is_empty_html = True

    if is_empty_html:
        try:
            from exporter.download_utils import process_fallback_downloads
            logger.info(f"Văn bản {vb.item_id} không có nội dung HTML. Thử tải file đính kèm...")
            download_path = process_fallback_downloads(vb.item_id, str(out_dir))
            if download_path:
                logger.info(f"Đã tạo file từ đính kèm: {download_path}")
                return Path(download_path)
        except Exception as e:
            logger.error(f"Fallback download failed: {e}")

    # 1. ATTEMPT HTMLDOCX FIRST (Perfect visual fidelity for tables and centering)
    if html_str and html_str.strip() and not is_massive_doc:
        try:
            from htmldocx import HtmlToDocx
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(13)
            doc.add_heading(vb.title or vb.doc_number or "Văn bản", level=1)
            
            # Clean HTML to prevent htmldocx crashes
            soup = BeautifulSoup(html_str, "html.parser")
            for tag in soup.find_all(['v:shape', 'o:p', 'script', 'style', 'iframe']):
                tag.decompose()
                
            # CRITICAL FIX: Convert CSS 'display: table' to HTML <table> for modern VBPL documents
            import re
            for div in soup.find_all('div'):
                style = div.get('style', '').lower()
                if 'display: table;' in style or 'display:table;' in style:
                    div.name = 'table'
                    div['width'] = '100%'
                    tr = soup.new_tag('tr')
                    cells = [c for c in div.find_all('div', recursive=False) if 'display: table-cell' in c.get('style', '').lower() or 'display:table-cell' in c.get('style', '').lower()]
                    if cells:
                        for c in cells:
                            c.name = 'td'
                            c_style = c.get('style', '').lower()
                            if 'width:' in c_style:
                                match = re.search(r'width:\s*([^;]+)', c_style)
                                if match:
                                    c['width'] = match.group(1).strip()
                            
                            is_center = 'text-align: center' in c_style or 'text-align:center' in c_style
                            for inner_div in c.find_all('div'):
                                inner_div.name = 'p'
                                if is_center:
                                    inner_div['style'] = inner_div.get('style', '') + '; text-align: center;'
                                    
                            tr.append(c)
                        div.clear()
                        div.append(tr)

            clean_html = str(soup)
            if not clean_html.strip().startswith("<"):
                clean_html = f"<div>{clean_html}</div>"
                
            parser = HtmlToDocx()
            parser.add_html_to_document(clean_html, doc)
            doc.save(str(file_path))
            logger.info(f"Đã lưu DOCX (HtmlToDocx fidelity): {file_path}")
            return file_path
        except Exception as e:
            logger.warning(f"HtmlToDocx failed ({e}), falling back to Pandoc...")
    elif is_massive_doc:
        logger.warning(f"Document is massive ({len(html_str)} chars). Skipping HtmlToDocx to prevent freezing.")

    # 2. FALLBACK TO PYPANDOC (Robust structure, but might lose layout tables)
    if html_str and html_str.strip():
        import pypandoc
        html_content = f"<html><head><meta charset='utf-8'></head><body><h1>{vb.title or vb.doc_number or 'Văn bản'}</h1>{html_str}</body></html>"
        try:
            try:
                pypandoc.convert_text(html_content, 'docx', format='html', outputfile=str(file_path))
            except OSError:
                pypandoc.download_pandoc(delete_installer=True)
                pypandoc.convert_text(html_content, 'docx', format='html', outputfile=str(file_path))
                
            # Post-process centering for Pandoc
            try:
                if not is_massive_doc:
                    soup = BeautifulSoup(html_content, 'html.parser')
                    centered_texts = set()
                    for tag in soup.find_all(True):
                        if tag.get('align', '').lower() == 'center' or 'text-align:center' in tag.get('style', '').lower().replace(' ', ''):
                            text = tag.get_text(strip=True)
                            if text: centered_texts.add(text)
                            for child in tag.find_all(['p', 'div']):
                                child_text = child.get_text(strip=True)
                                if child_text: centered_texts.add(child_text)
                                        
                    if centered_texts:
                        doc_obj = Document(str(file_path))
                        for p in doc_obj.paragraphs:
                            if p.text.strip() in centered_texts:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for table in doc_obj.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        if p.text.strip() in centered_texts:
                                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc_obj.save(str(file_path))
            except Exception:
                pass

            logger.info(f"Đã lưu DOCX (Pandoc fallback): {file_path}")
            return file_path
        except Exception as e:
            logger.error(f"Pandoc fallback also failed: {e}")

    # 3. FINAL FALLBACK: RAW TEXT
    logger.warning("Sử dụng fallback ghi chữ thô.")
    doc = Document()
    doc.add_heading(vb.title or vb.doc_number or "Văn bản", level=1)
    if getattr(vb, "articles", []):
        for article in vb.articles:
            doc.add_heading(article.get("title", ""), level=3)
            for line in article.get("content", "").splitlines():
                if line.strip(): doc.add_paragraph(line.strip())
    else:
        for line in getattr(vb, "full_text", "").splitlines():
            if line.strip(): doc.add_paragraph(line.strip())
    doc.save(str(file_path))
    logger.info(f"Đã lưu DOCX (Raw Text): {file_path}")
    return file_path