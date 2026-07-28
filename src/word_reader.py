from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
import tempfile
from pathlib import Path

from .text_normalization import normalize_with_mapping


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for part in iter(lambda: f.read(1024 * 1024), b""):
            h.update(part)
    return h.hexdigest()


def _docx_blocks(path: Path) -> list[dict]:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(path)
    blocks: list[dict] = []
    paragraph_index = table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            text = Paragraph(child, document).text
            if text.strip():
                blocks.append({"block_type": "paragraph", "text": text,
                    "paragraph_index": paragraph_index, "table_index": None,
                    "row_index": None, "cell_index": None})
            paragraph_index += 1
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            for row_i, row in enumerate(table.rows):
                for cell_i, cell in enumerate(row.cells):
                    text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                    if text:
                        blocks.append({"block_type": "table_cell", "text": text,
                            "paragraph_index": None, "table_index": table_index,
                            "row_index": row_i, "cell_index": cell_i})
            table_index += 1
    for section_i, section in enumerate(document.sections):
        for kind, container in (("header", section.header), ("footer", section.footer)):
            text = "\n".join(p.text for p in container.paragraphs if p.text.strip())
            if text:
                blocks.append({"block_type": kind, "text": text,
                    "paragraph_index": section_i, "table_index": None,
                    "row_index": None, "cell_index": None})
    return blocks


def _convert_doc(path: Path, cache_dir: Path) -> tuple[Path, str]:
    cache_dir.mkdir(parents=True, exist_ok=True)
    target = cache_dir / f"{sha256_file(path)}.docx"
    if target.exists():
        return target, "libreoffice_cached"
    if not shutil.which("libreoffice"):
        raise RuntimeError("LibreOffice is required to read legacy .doc files")
    with tempfile.TemporaryDirectory() as tmp:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmp, str(path)],
            capture_output=True, text=True, timeout=120)
        converted = next(Path(tmp).glob("*.docx"), None)
        if result.returncode or converted is None:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.strip()}")
        shutil.copy2(converted, target)
    return target, "libreoffice"


def extract_document(path: str | Path, document_id: str | None = None,
                     conversion_cache: str | Path = "cache/converted_doc") -> dict:
    path = Path(path).resolve()
    read_path, method = path, "python-docx"
    if path.suffix.lower() == ".doc":
        read_path, method = _convert_doc(path, Path(conversion_cache))
    blocks = _docx_blocks(read_path)
    normalized_parts: list[str] = []
    raw_parts: list[str] = []
    normalized_to_raw: list[int] = []
    normalized_cursor = raw_cursor = 0
    for i, block in enumerate(blocks):
        raw_value = block.pop("text")
        value, local_mapping = normalize_with_mapping(raw_value)
        if not value:
            continue
        start = normalized_cursor
        raw_start = raw_cursor
        if normalized_parts:
            normalized_to_raw.append(raw_start - 1)
        normalized_to_raw.extend(raw_start + offset for offset in local_mapping)
        block.update(block_id=f"b{i:05d}", text=value, start_char=start, end_char=start + len(value),
                     raw_start_char=raw_start, raw_end_char=raw_start + len(raw_value))
        normalized_parts.append(value)
        raw_parts.append(raw_value)
        normalized_cursor += len(value) + 1
        raw_cursor += len(raw_value) + 1
    blocks = [b for b in blocks if "start_char" in b]
    text = "\n".join(normalized_parts)
    return {
        "document_id": document_id or path.parent.name,
        "source_file": str(path), "source_sha256": sha256_file(path),
        "extraction_method": method, "raw_text": "\n".join(raw_parts), "normalized_text": text,
        "normalized_to_raw": normalized_to_raw, "blocks": blocks,
    }


def cache_document(document: dict, cache_dir: str | Path) -> Path:
    target = Path(cache_dir) / f"{document['document_id']}.json"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(document, ensure_ascii=False), encoding="utf-8")
    return target
