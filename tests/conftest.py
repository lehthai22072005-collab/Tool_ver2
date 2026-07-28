from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("QUYẾT ĐỊNH")
    doc.add_paragraph("Số: 12/QĐ-TTg")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Ngày ban hành"
    table.cell(0, 1).text = "01/02/2024"
    doc.save(path)
    return path
